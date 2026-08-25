#!/usr/bin/env python3
"""Generate the one-page course-map docx (intro video asset, 0_2).

One page, problem-first: each module = the problem you'll hit, and where the
course solves it. Regenerate with:  python3 slides_intro/mapa_kurzu.py
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ACCENT = RGBColor(0x1A, 0x73, 0xE8)   # Google blue
DARK = RGBColor(0x20, 0x24, 0x28)
GREY = RGBColor(0x5F, 0x63, 0x68)

ROWS = [
    ("M01 · Prvý agent",
     "Function-calling loop z minulého kurzu funguje — ale schému, dispatch aj slučku píšete ručne, pre každý projekt odznova.",
     "ADK to celé napíše za vás. Váš kód dostane mená: LlmAgent, Runner, Event, Session — a spustíme ho s jednou zmenenou linkou."),
    ("M02 · Nástroje",
     "Agent potrebuje viac než vaše funkcie: cudzie API, hotový tool server, pomoc iného agenta.",
     "Štyri príchute nástrojov na jednom príbehu (IT helpdesk) + čo s nástrojmi, ktoré vedia niečo zničiť."),
    ("M03 · Sessions a state",
     "Používateľ sa vráti zajtra. Čo si agent pamätá — a kde to vlastne je uložené?",
     "Konverzácia + vytiahnuté fakty; prefixy kľúčov rozhodujú, čo prežije a pre koho."),
    ("M04 · Výmena modelu",
     "Model zdražel alebo vypadol. Koľko práce je prejsť na iný?",
     "Jeden riadok. Ten istý agent beží na piatich modeloch od rôznych vendorov."),
    ("M05 · Workflow agenty",
     "Úloha má pevné kroky: za sebou, naraz, alebo dokola, kým výsledok nie je dobrý.",
     "Sequential, Parallel a Loop — pipeline bez toho, aby o poradí rozhodoval model."),
    ("M06 · Multi-agent",
     "Jeden agent na všetko prestáva stačiť.",
     "Dva vzory spolupráce: odovzdať konverzáciu vs. poradiť sa so špecialistom — a kedy multi-agent nechať tak."),
    ("M07 · Callbacky",
     "Potrebujete zasiahnuť do behu: zablokovať vstup, začierniť údaje, mocknúť nástroj v teste.",
     "Váš kód pred a po každom kroku agenta — guardrail, redakcia PII, mock."),
    ("M08 · Perzistencia a pamäť",
     "Reštart procesu = amnézia. A pamäť má fungovať aj naprieč konverzáciami.",
     "Databázové sessions (prežijú reštart) + MemoryService (spomienky naprieč konverzáciami)."),
    ("M09 · Evaluácia",
     "„Zdá sa, že to funguje“ nie je meranie.",
     "Testy agenta ako súbory, dve skóre a LLM ako sudca odpovede."),
    ("M10 · Nasadenie",
     "Agent žije v notebooku — zákazník ho tam nenájde.",
     "Z notebooku na HTTP službu; Docker, cloud a produkčný checklist."),
    ("M11–M14 · Časť 2 (samoštúdium)",
     "A čo navyše ponúka natívne Gemini?",
     "Grounding cez Google Search, thinking budgets, hlas naživo, protokol A2A — ochutnáme M11, zvyšok je samoštúdium."),
]


def set_cell_borders(cell, color="D9DDE3"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


doc = Document()
sec = doc.sections[0]
sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
sec.top_margin = sec.bottom_margin = Cm(1.2)
sec.left_margin = sec.right_margin = Cm(1.5)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(9.5)
style.paragraph_format.space_after = Pt(0)

title = doc.add_paragraph()
run = title.add_run("Google ADK — mapa kurzu")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = DARK
title.paragraph_format.space_after = Pt(2)

sub = doc.add_paragraph()
run = sub.add_run(
    "Priame pokračovanie kurzu Úvod do GenAI v Pythone. Každý modul rieši jeden problém, "
    "na ktorý pri stavbe agentov narazíte — keď viete, ktorý vás páli, viete, kam siahnuť."
)
run.font.size = Pt(10.5)
run.font.color.rgb = GREY
sub.paragraph_format.space_after = Pt(8)

table = doc.add_table(rows=0, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
widths = (Cm(3.6), Cm(7.2), Cm(7.2))

hdr = table.add_row()
for i, text in enumerate(("Modul", "Máte tento problém…", "…tu ho riešime")):
    cell = hdr.cells[i]
    cell.width = widths[i]
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = ACCENT
    set_cell_borders(cell, color="1A73E8")

for module, problem, solution in ROWS:
    row = table.add_row()
    for i, text in enumerate((module, problem, solution)):
        cell = row.cells[i]
        cell.width = widths[i]
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        if i == 0:
            r.font.bold = True
            r.font.color.rgb = DARK
        elif i == 1:
            r.font.italic = True
            r.font.color.rgb = GREY
        else:
            r.font.color.rgb = DARK
        set_cell_borders(cell)

foot = doc.add_paragraph()
foot.paragraph_format.space_before = Pt(10)
r = foot.add_run("Ponáhľate sa?  Rýchla cesta kurzom (~1 hodina): M01 → M02 → M05 → M11.")
r.font.bold = True
r.font.size = Pt(10)
r.font.color.rgb = DARK

foot2 = doc.add_paragraph()
r = foot2.add_run("Materiály: github.com/robertbarcik/ADK-tutorial · Google Drive priečinok v popise kurzu · videá po slovensky, notebooky po anglicky")
r.font.size = Pt(9)
r.font.color.rgb = GREY

out = Path(__file__).parent / "Google_ADK_mapa_kurzu.docx"
doc.save(out)
print(f"saved {out}")
